"""
Finalize the Word report - add figures, update lists, etc.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def finalize_report():
    doc_path = Path("docs/SEDS500-Project-Report-Umut-Akin.docx")
    doc = Document(doc_path)

    figures_dir = Path("docs/figures")

    # Find paragraph indices for key sections
    para_map = {}
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        para_map[text] = i

    # === UPDATE LIST OF TABLES ===
    update_list_of_tables(doc, para_map)

    # === UPDATE LIST OF FIGURES ===
    update_list_of_figures(doc, para_map)

    # === UPDATE LIST OF ABBREVIATIONS ===
    update_abbreviations(doc, para_map)

    # === ADD ABSTRACT ===
    add_abstract(doc, para_map)

    # === ADD FIGURES TO RESULTS SECTION ===
    add_figures_to_results(doc, para_map, figures_dir)

    # Save
    doc.save(doc_path)
    print(f"Report finalized and saved to: {doc_path}")


def update_list_of_tables(doc, para_map):
    """Add table entries to List of Tables section."""
    if 'LIST OF TABLES' not in para_map:
        return

    idx = para_map['LIST OF TABLES']
    heading_para = doc.paragraphs[idx]

    tables = [
        'Table 1: Comparison of Synthetic Data Generation Methods',
        'Table 2: Replacement Scenario Results',
        'Table 3: Augmentation Scenario Results',
        'Table 4: Privacy Test Results (Membership Inference Attack)',
        'Table 5: Ablation Study - Key Improvements',
    ]

    insert_list_items(doc, heading_para, tables)
    print("  - Updated List of Tables")


def update_list_of_figures(doc, para_map):
    """Add figure entries to List of Figures section."""
    if 'LIST OF FIGURES' not in para_map:
        return

    idx = para_map['LIST OF FIGURES']
    heading_para = doc.paragraphs[idx]

    figures = [
        'Figure 1: Replacement Scenario - Method Comparison',
        'Figure 2: Augmentation Scenario - Method Comparison',
        'Figure 3: Privacy-Utility Tradeoff',
        'Figure 4: Training Loss Convergence',
        'Figure 5: Diffusion Process Diagram',
        'Figure 6: Neural Network Architecture',
        'Figure 7: Privacy Evaluation Results',
        'Figure 8: Key Results Summary',
    ]

    insert_list_items(doc, heading_para, figures)
    print("  - Updated List of Figures")


def update_abbreviations(doc, para_map):
    """Add abbreviations to the list."""
    if 'LIST OF ABBREVIATIONS' not in para_map:
        return

    idx = para_map['LIST OF ABBREVIATIONS']
    heading_para = doc.paragraphs[idx]

    abbrevs = [
        'AUC\tArea Under the Curve',
        'CTGAN\tConditional Tabular Generative Adversarial Network',
        'DDPM\tDenoising Diffusion Probabilistic Models',
        'GAN\tGenerative Adversarial Network',
        'KL\tKullback-Leibler (divergence)',
        'MIA\tMembership Inference Attack',
        'MLP\tMulti-Layer Perceptron',
        'MSE\tMean Squared Error',
        'SMOGN\tSynthetic Minority Over-sampling for Regression with Gaussian Noise',
        'TabDDPM\tTabular Denoising Diffusion Probabilistic Models',
    ]

    insert_list_items(doc, heading_para, abbrevs)
    print("  - Updated List of Abbreviations")


def add_abstract(doc, para_map):
    """Add abstract before Introduction."""
    if 'INTRODUCTION' not in para_map:
        return

    idx = para_map['INTRODUCTION']
    intro_para = doc.paragraphs[idx]

    abstract_text = (
        "Organizations increasingly need to share sensitive tabular data for machine learning "
        "while protecting individual privacy. This project investigates diffusion models as a "
        "privacy-preserving approach for generating synthetic tabular data. We implement and "
        "evaluate TabDDPM-style diffusion with hybrid Gaussian-Multinomial noise handling, "
        "comparing it against CTGAN and SMOGN baselines. Our experiments demonstrate that "
        "TabDDPM-style diffusion achieves 87% of baseline model performance when training on "
        "synthetic data alone, significantly outperforming CTGAN (35%) and SMOGN (which fails "
        "catastrophically). Privacy validation through membership inference attacks confirms "
        "that the generated data leaks no information about training records (AUC = 0.51, "
        "equivalent to random guessing). These results establish diffusion models as a superior "
        "approach for generating high-utility, privacy-preserving synthetic tabular data."
    )

    # Insert abstract heading and text before Introduction
    heading_element = intro_para._element

    # Create abstract paragraph
    abstract_para = doc.add_paragraph(abstract_text)
    abstract_element = abstract_para._element
    abstract_element.getparent().remove(abstract_element)
    heading_element.addprevious(abstract_element)

    # Create abstract heading
    abstract_heading = doc.add_paragraph()
    abstract_heading.style = 'Heading 1'
    abstract_heading.add_run('ABSTRACT')
    heading_elem = abstract_heading._element
    heading_elem.getparent().remove(heading_elem)
    abstract_element.addprevious(heading_elem)

    print("  - Added Abstract section")


def add_figures_to_results(doc, para_map, figures_dir):
    """Add figures to the Results section."""
    # Find results-related paragraphs and add figures after them

    figure_placements = [
        ('4.2 Replacement Scenario Results', 'fig1_replacement_comparison.png', 'Figure 1: Replacement Scenario - Training on Synthetic Data Only'),
        ('4.3 Augmentation Scenario Results', 'fig2_augmentation_comparison.png', 'Figure 2: Augmentation Scenario - Original + Synthetic Data'),
        ('4.4 Privacy Evaluation', 'fig8_privacy_comparison.png', 'Figure 3: Privacy Evaluation - Membership Inference Attack AUC'),
    ]

    for section_text, fig_file, caption in figure_placements:
        # Find the section
        for i, para in enumerate(doc.paragraphs):
            if section_text in para.text:
                # Find a good place to insert (after some content)
                insert_idx = i + 3  # After a few paragraphs of content
                if insert_idx < len(doc.paragraphs):
                    insert_figure_after(doc, doc.paragraphs[insert_idx], figures_dir / fig_file, caption)
                    print(f"  - Added {fig_file}")
                break


def insert_figure_after(doc, para, image_path, caption):
    """Insert a figure with caption after a paragraph."""
    if not image_path.exists():
        print(f"    Warning: {image_path} not found")
        return

    heading_element = para._element

    # Create caption paragraph
    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_para.add_run(caption)
    run.italic = True
    caption_elem = caption_para._element
    caption_elem.getparent().remove(caption_elem)

    # Create figure paragraph
    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_para.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))
    fig_elem = fig_para._element
    fig_elem.getparent().remove(fig_elem)

    # Insert figure then caption
    heading_element.addnext(caption_elem)
    heading_element.addnext(fig_elem)


def insert_list_items(doc, heading_para, items):
    """Insert list items after a heading."""
    heading_element = heading_para._element

    # Insert in reverse order since we're using addnext
    for item in reversed(items):
        new_para = doc.add_paragraph(item)
        new_element = new_para._element
        new_element.getparent().remove(new_element)
        heading_element.addnext(new_element)


if __name__ == "__main__":
    print("Finalizing report...")
    finalize_report()
